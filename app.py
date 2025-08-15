# Import library yang dibutuhkan
import streamlit as st
import pandas as pd
import altair as alt
import google.generativeai as genai
from duckduckgo_search import DDGS
from googlesearch import search
import fitz  # PyMuPDF
import io
from docx import Document # Library baru untuk Word
import time

# ===================================================================
# KONFIGURASI HALAMAN
st.set_page_config(
    page_title="Policy Brief Generator Jateng",
    page_icon="📈",
    layout="wide"
)
# ===================================================================

# --- KONFIGURASI AI ---
try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    model = genai.GenerativeModel('gemini-1.5-flash')
except (KeyError, AttributeError):
    st.error("🚨 Kunci API Google tidak ditemukan! Pastikan file .streamlit/secrets.toml sudah benar.")
    st.stop()

# --- FUNGSI-FUNGSI UTAMA ---

def convert_to_docx(markdown_text):
    document = Document()
    # Membersihkan judul dari format markdown '**'
    title_raw = markdown_text.split('\n')[0]
    title_text = title_raw.replace('**Judul:**', '').strip()
    
    document.add_heading(title_text, level=1)
    
    # Mengambil sisa konten
    content_lines = markdown_text.split('\n')[1:]
    
    for line in content_lines:
        stripped_line = line.strip()
        if stripped_line.startswith('###'):
            heading_text = stripped_line.replace('###', '').strip()
            document.add_heading(heading_text, level=3)
        elif stripped_line.startswith('- '):
             # Menambahkan list bullet sederhana
            document.add_paragraph(stripped_line.replace('- ','').strip(), style='List Bullet')
        elif stripped_line:
            # Mengabaikan garis pemisah tabel markdown agar tidak tercetak
            if '|' in stripped_line and '---' in stripped_line:
                continue
            # Menambahkan paragraf biasa (termasuk baris tabel)
            document.add_paragraph(stripped_line)
            
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

# ... (Semua fungsi lain seperti extract_text_from_pdf, cari_dengan_duckduckgo, generate_brief_dengan_ai, dll. tetap sama) ...
def extract_text_from_pdf(pdf_file):
    st.write("📄 Mengekstrak teks dari file PDF...")
    try:
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        st.success("Ekstraksi teks dari PDF berhasil!")
        return text
    except Exception as e:
        st.error(f"Gagal memproses file PDF: {e}")
        return None
def cari_dengan_duckduckgo(keyword):
    st.write("🔎 Mencari dengan **DuckDuckGo**...")
    keyword_lokal = f'"{keyword} jawa tengah"'
    query = f'{keyword_lokal} (site:jatengprov.go.id OR site:undip.ac.id OR site:uns.ac.id OR site:unnes.ac.id OR site:walisongo.ac.id OR site:sinta.kemdikbud.go.id)'
    try:
        with DDGS() as ddgs:
            results = [r for r in ddgs.text(query, max_results=7)]
        if not results: return None, None
        snippets = [result['body'] for result in results]
        konteks = "\n\n".join(snippets)
        source_urls = [result['href'] for result in results]
        return konteks, source_urls
    except Exception as e:
        st.error(f"Error DuckDuckGo: {e}")
        return None, None
def cari_dengan_google(keyword):
    st.write("🔎 Mencari dengan **Google**...")
    keyword_lokal = f'"{keyword} jawa tengah"'
    query = f'{keyword_lokal} (site:jatengprov.go.id OR site:undip.ac.id OR site:uns.ac.id OR site:unnes.ac.id OR site:walisongo.ac.id OR site:sinta.kemdikbud.go.id)'
    try:
        search_results_urls = [url for url in search(query, num_results=10, sleep_interval=2)]
        if not search_results_urls: return None, None
        konteks = "Daftar sumber relevan:\n" + "\n".join(f"- {url}" for url in search_results_urls)
        return konteks, search_results_urls
    except Exception as e:
        st.error(f"Error Google: {e}")
        return None, None
def generate_brief_dengan_ai(sumber_info, konteks):
    st.write("🤖 AI menganalisis & menyusun draf teks...")
    prompt = f"""
    Anda adalah seorang analis kebijakan publik senior yang sangat teliti dan analitis.
    Tugas Anda adalah melakukan SINTESIS komprehensif dari informasi dalam 'KONTEKS' untuk menyusun draf policy brief yang kaya data dan mendalam.
    Sumber informasi untuk analisis ini adalah: {sumber_info}.
    INSTRUKSI PENTING:
    -   Gunakan HANYA informasi dari 'KONTEKS'.
    -   Jika menemukan data perbandingan atau statistik, formatlah dalam Tabel Markdown.
    -   Jangan membuat bagian referensi atau grafik, karena itu akan ditambahkan oleh sistem.
    KONTEKS:
    ---
    {konteks}
    ---
    Setelah menganalisis, tuliskan draf policy brief dengan format berikut (HANYA 4 bagian ini):
    **Judul:** [Judul yang tajam dan mencerminkan analisis]
    ### 1. Ringkasan Eksekutif
    [Ringkasan padat, sertakan 1-2 data kunci, dan sebutkan rekomendasi utama.]
    ### 2. Pendahuluan
    [Latar belakang masalah dengan data skala dan urgensi.]
    ### 3. Temuan dan Pembahasan Mendalam
    [Susun pembahasan berdasarkan TEMA UTAMA. Sajikan data, gunakan tabel jika cocok, bandingkan informasi, dan berikan analisis.]
    ### 4. Rekomendasi Kebijakan Berbasis Bukti (Evidence-Based)
    [Rekomendasi yang SPESIFIK, TERUKUR, dan dapat DILAKSANAKAN.]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error saat generate brief: {e}")
        return None
def generate_chart_code(konteks):
    st.write("📊 AI merancang kode untuk visualisasi data (fokus pada bar chart)...")
    prompt = f"""
    Anda adalah seorang spesialis visualisasi data menggunakan Python dan Altair.
    Tugas Anda adalah menganalisis teks 'KONTEKS' dan membuat satu buah kode Python sederhana untuk memvisualisasikan data yang paling penting dari teks tersebut.
    KONTEKS:
    ---
    {konteks}
    ---
    INSTRUKSI KODE YANG SANGAT KETAT:
    1.  Cari data numerik yang bisa dibandingkan. Prioritaskan data perbandingan antar kategori atau wilayah.
    2.  **Fokus utama adalah membuat diagram batang (bar chart) sederhana menggunakan `alt.Chart(data).mark_bar()`.** Jangan gunakan jenis chart lain kecuali sangat terpaksa.
    3.  Jika tidak ada data yang cocok untuk divisualisasikan, cukup kembalikan teks: #TIDAK_ADA_DATA
    4.  Pastikan kode lengkap, termasuk pembuatan `pandas.DataFrame`. Data dalam DataFrame harus berupa list sederhana.
    5.  Variabel final yang berisi chart **HARUS** bernama `chart`.
    6.  Jangan menulis `st.altair_chart(chart)` atau `print()`. Hanya kode pembuatan objek chart-nya saja.
    Contoh Sempurna:
    ```python
    import pandas as pd
    import altair as alt
    # Data sederhana dalam DataFrame
    data_source = pd.DataFrame({{'Kategori': ['UMKM Go Digital', 'UMKM Tradisional'],'Jumlah': [1500, 4500]}})
    # Pembuatan chart yang simpel dan pasti berhasil
    chart = alt.Chart(data_source).mark_bar().encode(
        x=alt.X('Kategori:N', title='Jenis UMKM', sort=None),
        y=alt.Y('Jumlah:Q', title='Jumlah UMKM'),
        tooltip=['Kategori', 'Jumlah']
    ).properties(
        title='Perbandingan Jumlah UMKM'
    )
    ```
    """
    try:
        response = model.generate_content(prompt)
        clean_code = response.text.replace("```python", "").replace("```", "").strip()
        return clean_code
    except Exception as e:
        st.error(f"Error saat generate chart code: {e}")
        return None

# --- TAMPILAN UTAMA APLIKASI ---
st.title("📄 Policy Brief Generator Pro")
st.write("Alat bantu AI untuk menyusun draf *policy brief* dari berbagai sumber.")

tab1, tab2 = st.tabs(["Cari dari Web 🌐", "Unggah Dokumen PDF 📄"])

# --- Logika untuk TAB 1: CARI DARI WEB ---
with tab1:
    st.header("Analisis Berdasarkan Pencarian Web")
    with st.form("search_form"):
        keyword_input = st.text_input("Masukkan Topik:", placeholder="Contoh: digitalisasi umkm di solo")
        search_engine = st.radio("Mesin Pencari:", ('DuckDuckGo', 'Google'), horizontal=True)
        submitted_search = st.form_submit_button("🚀 Buat Draf dari Web")

    if submitted_search and keyword_input:
        with st.spinner("Proses analisis komprehensif dari web..."):
            konteks, sumber_referensi = (None, None)
            if search_engine == 'DuckDuckGo':
                konteks, sumber_referensi = cari_dengan_duckduckgo(keyword_input)
            else:
                konteks, sumber_referensi = cari_dengan_google(keyword_input)

            if konteks and sumber_referensi:
                st.session_state.hasil_brief = generate_brief_dengan_ai(f"Pencarian web: '{keyword_input}'", konteks)
                st.session_state.kode_grafik = generate_chart_code(konteks)
                st.session_state.sumber_referensi = sumber_referensi
                st.session_state.keyword = keyword_input.replace(' ', '_')

# --- Logika untuk TAB 2: UNGGAH DOKUMEN PDF ---
with tab2:
    st.header("Analisis Berdasarkan Dokumen PDF")
    uploaded_file = st.file_uploader("Pilih file PDF", type="pdf")

    if uploaded_file:
        if st.button("🚀 Buat Draf dari PDF"):
            with st.spinner("Membaca & menganalisis PDF..."):
                konteks_pdf = extract_text_from_pdf(uploaded_file)
                if konteks_pdf:
                    st.session_state.hasil_brief = generate_brief_dengan_ai(f"Dokumen PDF: '{uploaded_file.name}'", konteks_pdf)
                    st.session_state.kode_grafik = generate_chart_code(konteks_pdf)
                    st.session_state.sumber_referensi = None # Tidak ada referensi URL untuk PDF
                    st.session_state.keyword = uploaded_file.name.split('.')[0].replace(' ', '_')


# --- BAGIAN TAMPILAN HASIL (DI LUAR TAB) ---
if 'hasil_brief' in st.session_state and st.session_state.hasil_brief:
    st.divider()
    st.header("Hasil Draf Policy Brief")

    # Tombol Unduh ditempatkan di atas agar mudah diakses
    docx_file = convert_to_docx(st.session_state.hasil_brief)
    st.download_button(
        label="📥 Unduh sebagai .docx",
        data=docx_file,
        file_name=f"policy_brief_{st.session_state.keyword}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.markdown(st.session_state.hasil_brief)

    if st.session_state.kode_grafik and "#TIDAK_ADA_DATA" not in st.session_state.kode_grafik:
        st.markdown("### Visualisasi Data Kunci")
        try:
            exec(st.session_state.kode_grafik, globals(), locals())
            st.altair_chart(locals()['chart'], use_container_width=True)
        except Exception as e:
            st.warning(f"Gagal membuat grafik: {e}")
            st.code(st.session_state.kode_grafik)

    if st.session_state.sumber_referensi:
        st.markdown("### Referensi")
        for url in st.session_state.sumber_referensi:
            st.markdown(f"- {url}")
    else:
        st.markdown("### Sumber Dokumen")

        st.info(f"Analisis ini dibuat berdasarkan dokumen yang diunggah.")
